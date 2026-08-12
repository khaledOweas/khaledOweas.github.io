#!/usr/bin/env python3
"""Generate ATS-friendly CV v13 (Word + PDF) for Khaled Ismail Oweas."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate

OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "13-Khaled Oweas CV.docx"
PDF_PATH = OUT_DIR / "13-Khaled Oweas CV.pdf"

NAME = "Khaled Ismail Oweas"
TITLE = "Senior / Lead Software Engineer"
EMAIL = "khaled.oweas@icloud.com"
PHONE = "+20 114 882 4462"
LOCATION = "Egypt · Remote (US timezones)"
LINKEDIN = "https://www.linkedin.com/in/khaledoweas"
GITHUB = "https://github.com/khaledOweas"
PORTFOLIO = "https://khaledoweas.github.io/"

SUMMARY = (
    "Senior/Lead software engineer with 10+ years designing, building, and operating "
    "production systems end to end — from database schema to deployed release. Strongest "
    "foundation in the .NET ecosystem (C#, ASP.NET Core, Web API, SQL Server, Entity "
    "Framework), with deep frontend experience in Angular and React with TypeScript. "
    "Based in Egypt and working remote-only with US companies across American timezones. "
    "Recent focus includes multi-tenant and multi-database architecture, application "
    "performance investigation, CI/CD with custom build agents, and dynamic reporting "
    "platforms. Independently designed and evolved live healthcare ecosystems used daily "
    "by patients, doctors, and staff (web, backend, database, and mobile apps on the App "
    "Store and Google Play)."
)

SKILLS = [
    (
        "Backend",
        "C#, VB.NET, ASP.NET Core, Web API, ASP.NET MVC, Razor, Web Forms, Blazor, "
        "EF Core, ADO.NET, LINQ, REST",
    ),
    (
        "Frontend",
        "React, Angular, AngularJS, TypeScript, JavaScript, jQuery, HTML/CSS, "
        "Bootstrap, Tailwind, Angular Material",
    ),
    (
        "Database",
        "SQL Server, MySQL, PostgreSQL, MongoDB, ElasticSearch, multi-tenant databases, "
        "stored procedures",
    ),
    (
        "Architecture",
        "Clean Architecture, CQRS, MediatR, microservices, multi-tenant systems, SOLID, "
        "design patterns",
    ),
    (
        "DevOps & Delivery",
        "CI/CD, custom build agents, Git, GitHub Actions, Docker, JIRA, cloud "
        "troubleshooting, production support",
    ),
    (
        "Mobile & Desktop",
        "Flutter, Dart, iOS, Android, Firebase, WPF, WinForms, legacy modernization",
    ),
    (
        "Real-time & Integrations",
        "SignalR, WebSockets, RabbitMQ, MSMQ, FCM, WhatsApp Cloud API, Kashier, "
        "Telegram, Facebook",
    ),
    (
        "Additional",
        "Python (scraping & automations), Arabic & English",
    ),
]

EXPERIENCE = [
    {
        "role": "Senior Software Engineer",
        "company": "Infinite Software Solutions, Inc. (ISSI) · USA · Remote",
        "dates": "Aug 2024 – Present",
        "bullets": [
            "Own large multi-tenant, multi-database enterprise systems for a US company while based in Egypt.",
            "Drive performance work that reduces heavy query and dashboard load times across API, database, and reporting layers.",
            "Build and maintain CI/CD pipelines with custom cloud/build agents; keep a steady release cadence with hands-on production troubleshooting.",
            "Deliver features on a dynamic reporting and analytics platform (React + TypeScript, .NET Core API).",
        ],
    },
    {
        "role": "Senior Software Engineer",
        "company": "Global Brands · Kaizen (Toshiba El Araby)",
        "dates": "~2022–2023",
        "bullets": [
            "Full-stack ownership of major features on Kaizen, a large industrial/manufacturing enterprise system.",
            "Mentored and trained two developers while carrying senior backend and frontend responsibility.",
        ],
    },
    {
        "role": "Senior .NET Consultant",
        "company": "Egabi Solutions (Egyptian Ministry of Justice)",
        "dates": "Nov 2022 – Mar 2023",
        "bullets": [
            "Built and supported government legal systems: Misdemeanors Court (.NET Web Forms) and Cassation Court (.NET Core + Angular).",
            "Designed scalable backends and SQL Server schemas, optimized stored procedures, and supported Legal portals (Cyclopedia, KOHA, Healthcare).",
        ],
    },
    {
        "role": "Backend / Microservices Engineer",
        "company": "ATIS, Albania · TaskWarp & NEMS",
        "dates": "Apr 2022 – Nov 2022",
        "bullets": [
            "Delivered .NET Core microservices on TaskWarp and .NET Core + Angular work on Nord Engineering / NEMS in an international team.",
        ],
    },
    {
        "role": "Senior .NET Developer",
        "company": "Mauve Mailorder Software, Germany",
        "dates": "~2022",
        "bullets": [
            "Designed and maintained backend services and a WPF CRM; owned MySQL modeling, optimization, integrations, and live-product stability.",
        ],
    },
    {
        "role": ".NET / IoT Engineer",
        "company": "ITE CORP, Egypt",
        "dates": "Apr 2022 – Sep 2022",
        "bullets": [
            "Shipped IoT Tracking for assets and people (.NET Framework) and an Innovation System (.NET Framework + AngularJS).",
        ],
    },
    {
        "role": ".NET Developer",
        "company": "DevopSolution / GAMA, KSA",
        "dates": "~2021 – Apr 2022",
        "bullets": [
            "Built Travotels travel APIs (.NET Core Web API) and Taif hospitality web apps (.NET Core MVC) for Saudi market operations.",
        ],
    },
    {
        "role": "Full-Stack / Backend Engineer",
        "company": "Alkhwadm Digital, KSA",
        "dates": "Mar 2021 – Jun 2021",
        "bullets": [
            "Shipped Call Center platform (AngularJS + .NET MVC) and WhatsApp Loop Flutter client for messaging-driven customer service.",
        ],
    },
    {
        "role": "R&D Software Engineer",
        "company": "eBSEG, Egypt",
        "dates": "Mar 2019 – Mar 2021",
        "bullets": [
            "R&D on omni-channel CX platforms (Portal, CEEP, Channel Manager, E-Messaging) and a ChatBot engine with full-text search / NLP.",
            "Integrated Facebook, WhatsApp, and Telegram for banking (NBE, CA, EBE, BSF, Bloom) and insurance (MLI, AHLIA) clients.",
        ],
    },
    {
        "role": "Full-Stack / Mobile Engineer",
        "company": "Intelligent Valley",
        "dates": "~2018–2019",
        "bullets": [
            "Shipped e-commerce web, My iFlights, Car Exhibitions, and Kunafa Shop across .NET backends, Angular, and Flutter mobile.",
        ],
    },
    {
        "role": ".NET / Angular Developer",
        "company": "Bayanatech for I.T., KSA",
        "dates": "~2017–2018",
        "bullets": [
            "Contributed to Tameer Pro construction-management platform (Angular + .NET Core APIs) for Saudi project workflows.",
        ],
    },
    {
        "role": "Founding Engineer",
        "company": "Medical Laboratory Management System (→ EKO) · Independent",
        "dates": "2015 – Present",
        "bullets": [
            "Started as desktop + web lab software (C# / VB.NET / WinForms / ASP.NET); modernized to Angular + .NET Core with a multi-app healthcare ecosystem for clients in Egypt, London, Saudi Arabia, and Libya.",
        ],
    },
]

PROJECTS = [
    {
        "name": "Integrated Laboratory Management & Digital Healthcare Platform (EKO)",
        "detail": (
            "Production healthcare ecosystem — web, backend, database, and role-specific "
            "mobile apps. Legacy WinForms/VB.NET/ASP.NET modernized to Angular + .NET Core; "
            "Blazor dispensing; WhatsApp Cloud API; EMR-oriented dashboards. Live in production."
        ),
    },
    {
        "name": "Khaled Dewan Imaging Center",
        "detail": (
            "End-to-end radiology booking, payments (Kashier), results delivery, Firebase/FCM "
            "and SignalR real-time updates, WhatsApp support. Live on App Store and Google Play (Flutter)."
        ),
    },
    {
        "name": "Enterprise Multi-Tenant Reporting & Analytics Platform (ISSI)",
        "detail": (
            "Dynamic reporting platform with multi-tenant/multi-database architecture, "
            "configurable dashboards, scheduled email jobs, and heavy query/dashboard performance focus "
            "(React + TypeScript, .NET Core API)."
        ),
    },
    {
        "name": "Additional selected deliveries",
        "detail": (
            "Clinic App; Website Builder CMS; OIL Desalter real-time CRUD; IoT Smart Home; "
            "Wagon Car Rent; Night-Tech API; Property Real Estate API; Call Center & WhatsApp Loop; "
            "IoT Tracking; Innovation System — .NET, Angular/AngularJS, Flutter."
        ),
    },
]

EDUCATION = {
    "degree": "Diploma of Technical Institutes",
    "school": "Industrial Institute technician · Technological College",
    "meta": "Section: Modern Electronics · Grade: Good · 2012 – 2014",
    "project": "Graduation project: Refrigerator and Freezer Temperature Monitoring",
}


def _xml(text: str) -> str:
    """Escape text for ReportLab Paragraph XML."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _tight_paragraph(paragraph, before=0, after=4, line=1.08):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def build_docx(path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, 0, 2, 1.0)
    r = p.add_run(NAME)
    _set_run_font(r, size=18, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, 0, 4, 1.0)
    r = p.add_run(TITLE)
    _set_run_font(r, size=11, bold=True, color=RGBColor(0x33, 0x33, 0x33))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, 0, 1, 1.05)
    r = p.add_run(f"{EMAIL}  ·  {PHONE}  ·  {LOCATION}")
    _set_run_font(r, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_paragraph(p, 0, 2, 1.05)
    r = p.add_run("LinkedIn: ")
    _set_run_font(r, size=9)
    _add_hyperlink(p, "linkedin.com/in/khaledoweas", LINKEDIN)
    r = p.add_run("  ·  GitHub: ")
    _set_run_font(r, size=9)
    _add_hyperlink(p, "github.com/khaledOweas", GITHUB)
    r = p.add_run("  ·  Full Profile: ")
    _set_run_font(r, size=9)
    _add_hyperlink(p, PORTFOLIO, PORTFOLIO)

    def section_heading(text: str):
        p = doc.add_paragraph()
        _tight_paragraph(p, 8, 2, 1.0)
        r = p.add_run(text.upper())
        _set_run_font(r, size=11, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "333333")
        pBdr.append(bottom)
        pPr.append(pBdr)

    section_heading("Professional Summary")
    p = doc.add_paragraph()
    _tight_paragraph(p, 2, 2, 1.08)
    r = p.add_run(SUMMARY)
    _set_run_font(r, size=10)

    section_heading("Core Skills")
    for label, items in SKILLS:
        p = doc.add_paragraph()
        _tight_paragraph(p, 1, 1, 1.05)
        r = p.add_run(f"{label}: ")
        _set_run_font(r, size=9.5, bold=True)
        r = p.add_run(items)
        _set_run_font(r, size=9.5)

    section_heading("Professional Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        _tight_paragraph(p, 5, 0, 1.05)
        r = p.add_run(job["role"])
        _set_run_font(r, size=10.5, bold=True)
        r = p.add_run(f"  |  {job['dates']}")
        _set_run_font(r, size=9.5, color=RGBColor(0x44, 0x44, 0x44))

        p = doc.add_paragraph()
        _tight_paragraph(p, 0, 1, 1.05)
        r = p.add_run(job["company"])
        _set_run_font(r, size=9.5, color=RGBColor(0x33, 0x33, 0x33))

        for bullet in job["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            _tight_paragraph(bp, 0, 1, 1.05)
            bp.clear()
            r = bp.add_run(bullet)
            _set_run_font(r, size=9.5)

    section_heading("Selected Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph()
        _tight_paragraph(p, 3, 0, 1.05)
        r = p.add_run(proj["name"])
        _set_run_font(r, size=10, bold=True)
        p = doc.add_paragraph()
        _tight_paragraph(p, 0, 2, 1.05)
        r = p.add_run(proj["detail"])
        _set_run_font(r, size=9.5)

    section_heading("Education")
    p = doc.add_paragraph()
    _tight_paragraph(p, 3, 0, 1.05)
    r = p.add_run(EDUCATION["degree"])
    _set_run_font(r, size=10.5, bold=True)
    p = doc.add_paragraph()
    _tight_paragraph(p, 0, 0, 1.05)
    r = p.add_run(EDUCATION["school"])
    _set_run_font(r, size=9.5)
    p = doc.add_paragraph()
    _tight_paragraph(p, 0, 0, 1.05)
    r = p.add_run(EDUCATION["meta"])
    _set_run_font(r, size=9.5)
    p = doc.add_paragraph()
    _tight_paragraph(p, 0, 0, 1.05)
    r = p.add_run(EDUCATION["project"])
    _set_run_font(r, size=9.5)

    doc.save(path)


def build_pdf(path: Path) -> None:
    accent = HexColor("#1a1a1a")
    muted = HexColor("#444444")

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        title=f"{NAME} — CV",
        author=NAME,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CVName",
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            alignment=TA_CENTER,
            textColor=accent,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVTitle",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            alignment=TA_CENTER,
            textColor=muted,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVContact",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
            textColor=black,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVHeading",
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=accent,
            spaceBefore=8,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBody",
            fontName="Helvetica",
            fontSize=9,
            leading=11.5,
            alignment=TA_JUSTIFY,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVSkill",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVJobRole",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=12,
            spaceBefore=5,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVCompany",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            textColor=muted,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVBullet",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            leftIndent=10,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVProjName",
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=11.5,
            spaceBefore=3,
            spaceAfter=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CVProjDetail",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            spaceAfter=2,
        )
    )

    story = []
    story.append(Paragraph(_xml(NAME), styles["CVName"]))
    story.append(Paragraph(_xml(TITLE), styles["CVTitle"]))
    story.append(
        Paragraph(
            _xml(f"{EMAIL}  ·  {PHONE}  ·  {LOCATION}"),
            styles["CVContact"],
        )
    )
    story.append(
        Paragraph(
            f'LinkedIn: <link href="{LINKEDIN}" color="blue"><u>linkedin.com/in/khaledoweas</u></link>'
            f'  ·  GitHub: <link href="{GITHUB}" color="blue"><u>github.com/khaledOweas</u></link>'
            f'  ·  Full Profile: <link href="{PORTFOLIO}" color="blue"><u>{_xml(PORTFOLIO)}</u></link>',
            styles["CVContact"],
        )
    )

    def heading(text: str):
        story.append(Paragraph(_xml(text.upper()), styles["CVHeading"]))
        story.append(
            HRFlowable(
                width="100%",
                thickness=0.8,
                color=accent,
                spaceBefore=0,
                spaceAfter=4,
            )
        )

    heading("Professional Summary")
    story.append(Paragraph(_xml(SUMMARY), styles["CVBody"]))

    heading("Core Skills")
    for label, items in SKILLS:
        story.append(
            Paragraph(f"<b>{_xml(label)}:</b> {_xml(items)}", styles["CVSkill"])
        )

    heading("Professional Experience")
    for job in EXPERIENCE:
        story.append(
            Paragraph(
                _xml(f'{job["role"]}  |  {job["dates"]}'),
                styles["CVJobRole"],
            )
        )
        story.append(Paragraph(_xml(job["company"]), styles["CVCompany"]))
        for bullet in job["bullets"]:
            story.append(Paragraph(_xml(f"- {bullet}"), styles["CVBullet"]))

    heading("Selected Projects")
    for proj in PROJECTS:
        story.append(Paragraph(_xml(proj["name"]), styles["CVProjName"]))
        story.append(Paragraph(_xml(proj["detail"]), styles["CVProjDetail"]))

    heading("Education")
    story.append(Paragraph(_xml(EDUCATION["degree"]), styles["CVJobRole"]))
    story.append(Paragraph(_xml(EDUCATION["school"]), styles["CVCompany"]))
    story.append(Paragraph(_xml(EDUCATION["meta"]), styles["CVProjDetail"]))
    story.append(Paragraph(_xml(EDUCATION["project"]), styles["CVProjDetail"]))

    doc.build(story)


def main() -> None:
    build_docx(DOCX_PATH)
    build_pdf(PDF_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()
